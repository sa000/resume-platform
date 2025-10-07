"""
Text extraction utilities for resume parsing.

Supports PDF and DOCX formats, extracting content from both
standard paragraphs and embedded tables.
"""

import logging
import docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file.

    Extracts text from both regular paragraphs and table cells,
    preserving the document structure.

    Args:
        file_path: Path to the DOCX file

    Returns:
        str: Concatenated text content with newline separators
    """
    doc = docx.Document(file_path)
    text_chunks = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_chunks.append(paragraph.text)

    # Extract table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_chunks.append(cell_text)

    text = "\n".join(text_chunks)
    logger.debug(f"Extracted {len(text)} characters from DOCX")
    return text


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file.

    Args:
        file_path: Path to the PDF file

    Returns:
        str: Concatenated text from all pages
    """
    reader = PdfReader(file_path)
    text = "\n".join([page.extract_text() for page in reader.pages])
    logger.debug(f"Extracted {len(text)} characters from PDF")
    return text


def extract_text(file_path: str) -> str:
    """
    Extract text from a resume file based on file extension.

    Automatically detects file type and uses the appropriate
    extraction method.

    Args:
        file_path: Path to resume file (.pdf or .docx)

    Returns:
        str: Extracted text content

    Raises:
        ValueError: If file extension is not supported
    """
    if file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)
    elif file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")