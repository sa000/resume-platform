"""
Talent Search Platform - Streamlit Application

Main web application for searching and filtering hedge fund analyst candidates.
Provides full-text search, multi-dimensional filtering, analytics visualizations,
and detailed candidate profile views with resume preview.
"""

import streamlit as st
import pandas as pd
import sqlite3
import plotly.express as px
import pathlib
import json
import base64
import logging

from docx import Document

logger = logging.getLogger(__name__)

# =============================================================================
# CONFIGURATION
# =============================================================================
st.set_page_config(
    page_title="Talent Search Platform",
    page_icon="",
    layout="wide",
    initial_sidebar_state="collapsed"
)

DB_PATH = pathlib.Path(__file__).parents[1] / "data/db/warehouse.db"

# Load custom CSS styling
with open(pathlib.Path(__file__).parent / "styles.css") as f:
    st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


# =============================================================================
# DATA LOADING FUNCTIONS
# =============================================================================

def get_filter_values(field_name: str) -> list:
    """
    Get unique filter values from pre-computed lookup table.

    Uses indexed filter_values table for fast dropdown population.

    Args:
        field_name: Filter category (e.g., 'skill', 'company', 'geography')

    Returns:
        list: Sorted unique values
    """
    conn = sqlite3.connect(DB_PATH)
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT DISTINCT field_value
            FROM filter_values
            WHERE field_name = ?
            ORDER BY field_value
        """, (field_name,))

        values = [row[0] for row in cur.fetchall()]
        conn.close()
        return values
    except Exception as e:
        logger.error(f"Failed to get filter values for {field_name}: {e}")
        conn.close()
        return []


def load_candidates() -> pd.DataFrame:
    """
    Load all candidates with aggregated data.

    Performs joins across candidates, skills, experiences, education,
    and quality_scores tables to create a complete candidate view.

    Returns:
        pd.DataFrame: Candidates with all associated data
    """
    conn = sqlite3.connect(DB_PATH)
    df = pd.read_sql_query("""
        SELECT
            c.*,
            qs.quality_score,
            qs.grade AS quality_grade,
            qs.total_issues,
            GROUP_CONCAT(DISTINCT s.skill) AS all_skills,
            GROUP_CONCAT(DISTINCT e.company) AS all_companies,
            GROUP_CONCAT(DISTINCT ed.school) AS all_schools,
            GROUP_CONCAT(DISTINCT ed.degree) AS all_degrees
        FROM candidates c
        LEFT JOIN skills s ON s.candidate_id = c.id
        LEFT JOIN experiences e ON e.candidate_id = c.id
        LEFT JOIN education ed ON ed.candidate_id = c.id
        LEFT JOIN quality_scores qs ON qs.candidate_id = c.id
        GROUP BY c.id

    """, conn)
    conn.close()

    for col in ["top_skills", "notable_experience"]:
        if col in df.columns:
            df[col] = df[col].apply(lambda x: json.loads(x) if isinstance(x, str) and x.startswith("[") else [])

    # Determine highest degree for each candidate
    def get_highest_degree(degrees_str):
        if not degrees_str:
            return None
        degrees = str(degrees_str).upper().split(',')
        degree_hierarchy = {'PH.D.': 4, 'PHD': 4, 'PH.D': 4, 'MBA': 3, 'M.S.': 2, 'MS': 2, 'M.S': 2, 'B.S.': 1, 'BS': 1, 'B.S': 1, 'BA': 1, 'B.A.': 1}
        highest = None
        highest_rank = 0
        for degree in degrees:
            degree = degree.strip()
            for key, rank in degree_hierarchy.items():
                if key in degree:
                    if rank > highest_rank:
                        highest_rank = rank
                        if rank == 4:
                            highest = 'PhD'
                        elif rank == 3:
                            highest = 'MBA'
                        elif rank == 2:
                            highest = 'MS'
                        elif rank == 1:
                            highest = 'BS'
        return highest

    df['highest_degree'] = df['all_degrees'].apply(get_highest_degree)
    return df


def search_candidates(search_query: str) -> pd.DataFrame:
    """
    Full-text search using FTS5 with BM25 ranking.

    Searches across names, titles, companies, skills, education, and
    certifications. Returns results ranked by relevance.

    Args:
        search_query: Search terms (supports Boolean operators, phrases, wildcards)

    Returns:
        pd.DataFrame: Matching candidates with match_info column
        None: If search query is empty or search fails
    """
    if not search_query or search_query.strip() == "":
        return None  # Return None to indicate no search performed

    conn = sqlite3.connect(DB_PATH)

    # FTS5 query with ranking - join back to get all candidate data
    query = """
        SELECT
            c.*,
            qs.quality_score,
            qs.grade AS quality_grade,
            qs.total_issues,
            GROUP_CONCAT(DISTINCT s.skill) AS all_skills,
            GROUP_CONCAT(DISTINCT e.company) AS all_companies,
            GROUP_CONCAT(DISTINCT ed.school) AS all_schools,
            GROUP_CONCAT(DISTINCT ed.degree) AS all_degrees,
            fts.rank,
            fts.name as fts_name,
            fts.current_title as fts_title,
            fts.current_company as fts_company,
            fts.skills as fts_skills,
            fts.experience_text as fts_experience,
            fts.education_text as fts_education,
            fts.certifications as fts_certs
        FROM candidates_fts fts
        JOIN candidates c ON c.id = fts.candidate_id
        LEFT JOIN skills s ON s.candidate_id = c.id
        LEFT JOIN experiences e ON e.candidate_id = c.id
        LEFT JOIN education ed ON ed.candidate_id = c.id
        LEFT JOIN quality_scores qs ON qs.candidate_id = c.id
        WHERE candidates_fts MATCH ?
        GROUP BY c.id
        ORDER BY fts.rank
    """

    try:
        df = pd.read_sql_query(query, conn, params=(search_query,))
        conn.close()

        # Process JSON fields like in load_candidates
        for col in ["top_skills", "notable_experience"]:
            if col in df.columns:
                df[col] = df[col].apply(lambda x: json.loads(x) if isinstance(x, str) and x.startswith("[") else [])

        # Determine highest degree for each candidate (same as load_candidates)
        def get_highest_degree(degrees_str):
            if not degrees_str:
                return None
            degrees = str(degrees_str).upper().split(',')
            degree_hierarchy = {'PH.D.': 4, 'PHD': 4, 'PH.D': 4, 'MBA': 3, 'M.S.': 2, 'MS': 2, 'M.S': 2, 'B.S.': 1, 'BS': 1, 'B.S': 1, 'BA': 1, 'B.A.': 1}
            highest = None
            highest_rank = 0
            for degree in degrees:
                degree = degree.strip()
                for key, rank in degree_hierarchy.items():
                    if key in degree:
                        if rank > highest_rank:
                            highest_rank = rank
                            if rank == 4:
                                highest = 'PhD'
                            elif rank == 3:
                                highest = 'MBA'
                            elif rank == 2:
                                highest = 'MS'
                            elif rank == 1:
                                highest = 'BS'
            return highest

        df['highest_degree'] = df['all_degrees'].apply(get_highest_degree)

        # Determine what matched for each candidate
        def get_match_info(row):
            query_lower = search_query.lower()
            matches = []

            # Check each field for matches
            if row.get('fts_name') and query_lower in str(row['fts_name']).lower():
                matches.append(f" Name: {row['name']}")

            if row.get('fts_company') and query_lower in str(row['fts_company']).lower():
                matches.append(f" Company: {row['current_company']}")

            if row.get('fts_title') and query_lower in str(row['fts_title']).lower():
                matches.append(f" Title: {row['current_title']}")

            if row.get('fts_skills') and query_lower in str(row['fts_skills']).lower():
                # Find which skills matched
                skills = str(row.get('all_skills', '')).split(',')
                matched_skills = [s.strip() for s in skills if query_lower in s.lower()]
                if matched_skills:
                    matches.append(f" Skills: {', '.join(matched_skills[:3])}")

            if row.get('all_companies') and query_lower in str(row['all_companies']).lower():
                companies = str(row['all_companies']).split(',')
                matched_companies = [c.strip() for c in companies if query_lower in c.lower()]
                if matched_companies and matched_companies[0] != row.get('current_company'):
                    matches.append(f" Past Experience: {', '.join(matched_companies[:2])}")

            if row.get('fts_education') and query_lower in str(row['fts_education']).lower():
                matches.append(f" Education matched")

            if row.get('fts_certs') and query_lower in str(row['fts_certs']).lower():
                matches.append(f" Certifications matched")

            return matches if matches else [" Relevant match found"]

        df['match_info'] = df.apply(get_match_info, axis=1)

        return df
    except Exception as e:
        logger.error(f"Search failed: {e}")
        st.error(f"Search error: {e}")
        conn.close()
        return None


def get_search_suggestions() -> list:
    """
    Get common search terms for autocomplete.

    Returns:
        list: Sorted list combining top companies, skills, and degrees
    """
    conn = sqlite3.connect(DB_PATH)
    suggestions = []

    try:
        # Top companies
        companies = pd.read_sql_query(
            "SELECT DISTINCT company FROM experiences WHERE company IS NOT NULL ORDER BY company LIMIT 30",
            conn
        )
        suggestions.extend(companies['company'].tolist())

        # Top skills
        skills = pd.read_sql_query(
            "SELECT DISTINCT skill FROM skills WHERE skill IS NOT NULL ORDER BY skill LIMIT 30",
            conn
        )
        suggestions.extend(skills['skill'].tolist())

        # Degrees
        degrees = pd.read_sql_query(
            "SELECT DISTINCT degree FROM education WHERE degree IS NOT NULL ORDER BY degree",
            conn
        )
        suggestions.extend(degrees['degree'].tolist())

    except Exception as e:
        logger.error(f"Failed to get suggestions: {e}")

    conn.close()
    return sorted(set(suggestions))  # Remove duplicates and sort


# =============================================================================
# UI SETUP
# =============================================================================
st.markdown("""
    <div class="main-header">
        <h1> Talent Search Platform</h1>
        <p>Find the perfect candidates for your team with advanced filtering and search</p>
    </div>
""", unsafe_allow_html=True)

try:
    df = load_candidates()
    all_skills = get_filter_values("skill")
    all_companies = get_filter_values("company")
    all_schools = get_filter_values("school")
    all_degrees = get_filter_values("degree")
except Exception as e:
    st.error(f" Could not load data from database: {e}")
    st.stop()

if df.empty:
    st.warning(" No candidate data found in warehouse.")
    st.stop()

# =============================================================================
# SESSION STATE
# =============================================================================
AVAILABLE_ROLES = ["Role 1", "Role 2", "Role 3"]

if 'flagged_candidates' not in st.session_state:
    st.session_state.flagged_candidates = {}  # {candidate_id: [list of roles]}


# =============================================================================
# TAB NAVIGATION
# =============================================================================
tab1, tab2 = st.tabs([" Talent Search", " Work in Progress: Matching Jobs to Candidates"])

# =============================================================================
# TAB 2: JOB MATCHING (WORK IN PROGRESS)
# =============================================================================
with tab2:
    st.markdown('<h2 class="section-header"> Job-to-Candidate Matching (Work in Progress)</h2>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("###  Select a Job Role to identify the best Candidate for the Role")

    col1, col2 = st.columns([1, 3])

    with col1:
        if st.button(" Data Scientist: Hiring Manager Sammie Kim", type="primary", use_container_width=True):
            st.session_state.selected_job = "Data Scientist: HM Sammie"

    if st.session_state.get('selected_job'):
        st.markdown("---")
        st.markdown("###  Match Results")

        # Joking message about the match
        st.success("""
        ** Best Candidate Found: Sakibul Alam**

        *Congratulations! Our highly advanced AI algorithm has determined that
        Sakibul Alam is the perfect match for this role!* 
        """)

        st.markdown("---")
        st.markdown("###  Future Improvements")

        st.markdown("""
        Jokes aside a feature should be users should be able to upload a job description a link to the job posting and we 
        will present the most relevant candidates to them.

        We can do so by using AI to similarly extract relevant properties from the Job Posting like skills, experience, etc. Then do matching against our
        candidates.

        We can then do smarter methods like vector embeddings, similarities, feedback looks to imrprove the matching process
        """)

        st.markdown("---")

# =============================================================================
# TAB 1: TALENT SEARCH
# =============================================================================
with tab1:
    # Sidebar: Flagged candidates summary
    with st.sidebar:
        st.markdown("###  Flagged Candidates by Role")

        # Count candidates per role
        role_counts = {role: 0 for role in AVAILABLE_ROLES}
        for candidate_id, roles in st.session_state.flagged_candidates.items():
            for role in roles:
                if role in role_counts:
                    role_counts[role] += 1

        # Display counts
        for role in AVAILABLE_ROLES:
            st.metric(role, role_counts[role])

        st.markdown("---")
        total_flagged = len(st.session_state.flagged_candidates)
        st.metric("Total Candidates Flagged", total_flagged)

    # Flagged candidates review section
    if st.session_state.flagged_candidates:
        with st.expander(" Review Flagged Candidates", expanded=False):
            # Create tabs for each role
            tabs = st.tabs(AVAILABLE_ROLES + ["All Flagged"])

            for i, role in enumerate(AVAILABLE_ROLES):
                with tabs[i]:
                    # Get candidates flagged for this role
                    role_candidate_ids = [cid for cid, roles in st.session_state.flagged_candidates.items() if role in roles]

                    if role_candidate_ids:
                        role_candidates = df[df['id'].isin(role_candidate_ids)]

                        for _, cand in role_candidates.iterrows():
                            col1, col2, col3 = st.columns([3, 1, 1])
                            with col1:
                                st.markdown(f"**{cand['name']}** - {cand['current_title']} at {cand['current_company']}")
                                # Show all roles this candidate is flagged for
                                flagged_roles = st.session_state.flagged_candidates.get(cand['id'], [])
                                roles_display = ", ".join(flagged_roles)
                                st.caption(f"Flagged for: {roles_display}")
                            with col2:
                                st.caption(f"{cand['years_experience']} years exp")
                            with col3:
                                if st.button("Remove", key=f"remove_{role}_{cand['id']}"):
                                    st.session_state.flagged_candidates[cand['id']].remove(role)
                                    if not st.session_state.flagged_candidates[cand['id']]:
                                        del st.session_state.flagged_candidates[cand['id']]
                                    st.rerun()
                    else:
                        st.info(f"No candidates flagged for {role} yet.")

            # All flagged tab
            with tabs[-1]:
                if st.session_state.flagged_candidates:
                    all_flagged = df[df['id'].isin(st.session_state.flagged_candidates.keys())]

                    for _, cand in all_flagged.iterrows():
                        col1, col2, col3 = st.columns([3, 1, 1])
                        with col1:
                            st.markdown(f"**{cand['name']}** - {cand['current_title']} at {cand['current_company']}")
                            flagged_roles = st.session_state.flagged_candidates.get(cand['id'], [])
                            roles_display = ", ".join(flagged_roles)
                            st.caption(f"Flagged for: {roles_display}")
                        with col2:
                            st.caption(f"{cand['years_experience']} years exp")
                        with col3:
                            if st.button("Clear All", key=f"clear_all_{cand['id']}"):
                                del st.session_state.flagged_candidates[cand['id']]
                                st.rerun()
                else:
                    st.info("No flagged candidates yet.")


# =============================
    #  VISUALIZATIONS - MOVED TO TOP
    # =============================
    st.markdown('<h2 class="section-header"> Analytics & Insights</h2>', unsafe_allow_html=True)

    col_a, col_b = st.columns(2, gap="large")

    with col_a:
        st.markdown("####  Candidates by Geography")
        geo_counts = df["primary_geography"].value_counts().reset_index()
        geo_counts.columns = ["Geography", "Candidates"]
        fig_geo = px.bar(
            geo_counts,
            x="Geography",
            y="Candidates",
            text="Candidates",
            color_discrete_sequence=['#3b82f6']
        )
        fig_geo.update_layout(
            showlegend=False,
            template='plotly_dark',
            plot_bgcolor='#1e293b',
            paper_bgcolor='#1e293b',
            yaxis_title="Number of Candidates",
            xaxis=dict(showgrid=False),
            yaxis=dict(showgrid=False),
            font=dict(size=13),
            margin=dict(l=50, r=50, t=50, b=50),
            height=450
        )
        st.plotly_chart(fig_geo, use_container_width=True)

    with col_b:
        st.markdown("####  Candidates by Sector")
        sector_counts = df["primary_sector"].value_counts().reset_index()
        sector_counts.columns = ["Sector", "Candidates"]
        total_candidates = len(df)

        fig_sector = px.pie(
            sector_counts,
            names="Sector",
            values="Candidates",
            hole=0.5,
            color_discrete_sequence=px.colors.qualitative.Bold
        )
        fig_sector.update_layout(
            template='plotly_dark',
            plot_bgcolor='#1e293b',
            paper_bgcolor='#1e293b',
            font=dict(size=13),
            margin=dict(l=50, r=50, t=50, b=50),
            height=450,
            annotations=[dict(
                text=f'<b>{total_candidates}</b><br>Total',
                x=0.5, y=0.5,
                font_size=24,
                showarrow=False,
                font=dict(color='white')
            )]
        )
        st.plotly_chart(fig_sector, use_container_width=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    st.markdown("####  Candidates by Investment Approach")
    approach_counts = df["investment_approach"].value_counts().reset_index()
    approach_counts.columns = ["Approach", "Candidates"]
    fig_approach = px.bar(
        approach_counts,
        x="Approach",
        y="Candidates",
        text="Candidates",
        color_discrete_sequence=['#8b5cf6']
    )
    fig_approach.update_layout(
        showlegend=False,
        template='plotly_dark',
        plot_bgcolor='#1e293b',
        paper_bgcolor='#1e293b',
        yaxis_title="Number of Candidates",
        xaxis=dict(showgrid=False),
        yaxis=dict(showgrid=False),
        font=dict(size=13),
        margin=dict(l=50, r=50, t=50, b=50),
        height=450
    )
    st.plotly_chart(fig_approach, use_container_width=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Search and filters
    st.markdown('<h2 class="section-header"> Search & Filter Candidates</h2>', unsafe_allow_html=True)

    # Initialize search in session state
    if 'search_query' not in st.session_state:
        st.session_state.search_query = ""

    # Search bar at top of filters
    st.markdown("####  Smart Search")
    col_search1, col_search2, col_search3 = st.columns([5, 1, 1])

    with col_search1:
        search_input = st.text_input(
            "Search across all candidate data",
            value=st.session_state.search_query,
            placeholder="Try searching machine learning, then hit search",
            help="Live search across names, companies, skills, education, certifications, and experience",
            label_visibility="collapsed",
            key="search_input_field"
        )

    with col_search2:
        if st.button(" Search", type="primary", width='stretch'):
            st.session_state.search_query = search_input
            st.rerun()

    with col_search3:
        if st.button(" Clear", type="secondary", width='stretch'):
            st.session_state.search_query = ""
            st.rerun()

    # Use the session state search query
    search_query = st.session_state.search_query

    st.markdown("####  Additional Filters")
    col1, col2, col3, col4, col5, col6 = st.columns(6)

    with col1:
        geo = st.selectbox(" Geographic Market", ["All"] + sorted(df["primary_geography"].dropna().unique().tolist()))

    with col2:
        approach = st.selectbox(" Investment Approach", ["All"] + sorted(df["investment_approach"].dropna().unique().tolist()))

    with col3:
        sector = st.selectbox(" Sector", ["All"] + sorted(df["primary_sector"].dropna().unique().tolist()))

    with col4:
        education_degree = st.selectbox(" Education", ["All"] + all_degrees)

    with col5:
        company = st.selectbox(" Company (Any Past Employer)", ["All"] + all_companies)

    with col6:
        school = st.selectbox(" School", ["All"] + all_schools)

    st.markdown("####  Experience & Skills")
    col7, col8 = st.columns([1, 2])

    with col7:
        min_exp, max_exp = st.slider(
            "Years of Experience",
            min_value=0,
            max_value=int(df["years_experience"].max() or 20),
            value=(0, int(df["years_experience"].max() or 20)),
        )

    with col8:
        selected_skills = st.multiselect(" Skills", options=all_skills, placeholder="Choose skills...")


    # Apply search and filters
    if search_query and search_query.strip():
        search_results = search_candidates(search_query)

        if search_results is not None and not search_results.empty:
            filtered = search_results.copy()
            st.success(f" Found **{len(filtered)}** candidates matching: **{search_query}**")
        else:
            st.warning(" No candidates found matching your search query. Try different keywords.")
            filtered = pd.DataFrame()  # Empty dataframe
    else:
        # No search query - start with all candidates
        filtered = df.copy()

    # Apply additional filters on top of search results (if any)
    if not filtered.empty:
        # Filter by geography
        if geo != "All":
            filtered = filtered[filtered["primary_geography"] == geo]

        # Filter by investment approach
        if approach != "All":
            filtered = filtered[filtered["investment_approach"] == approach]

        # Filter by sector
        if sector != "All":
            filtered = filtered[filtered["primary_sector"] == sector]

        # Filter by education degree
        if education_degree != "All":
            filtered = filtered[
                filtered["all_degrees"].apply(lambda x: education_degree in str(x) if x else False)
            ]

        # Filter by experience range
        filtered = filtered[(filtered["years_experience"] >= min_exp) & (filtered["years_experience"] <= max_exp)]

        # Filter by company
        if company != "All":
            filtered = filtered[
                filtered["all_companies"].apply(lambda x: company.lower() in str(x).lower() if x else False)
            ]

        # Filter by school
        if school != "All":
            filtered = filtered[
                filtered["all_schools"].apply(lambda x: school.lower() in str(x).lower() if x else False)
            ]

        # Match any selected skill
        if selected_skills:
            filtered = filtered[
                filtered["all_skills"].apply(lambda x: any(skill.lower() in str(x).lower() for skill in selected_skills))
            ]


    # Results summary table
    st.markdown("<br>", unsafe_allow_html=True)
    if search_query and search_query.strip():
        st.markdown(f'<h2 class="section-header"> Search Results: {len(filtered)} Candidates</h2>', unsafe_allow_html=True)
        if len(filtered) > 0:
            st.caption(f"Showing results for '{search_query}' with applied filters")
    else:
        st.markdown(f'<h2 class="section-header"> All Candidates: {len(filtered)} Found</h2>', unsafe_allow_html=True)
    if len(filtered) == 0:
        st.warning(" No candidates match your criteria. Try adjusting filters or search terms.")
    elif len(filtered) < 5:
        st.success(f" Found {len(filtered)} highly relevant candidate(s)")
    else:
        st.info(f" Showing {len(filtered)} matching candidates")

    cols = [
        "name",
        "current_title",
        "current_company",
        "primary_sector",
        "investment_approach",
        "primary_geography",
        "years_experience",
    ]
    display_df = filtered[cols].copy()
    display_df.columns = ["Name", "Current Title", "Company", "Sector", "Investment Approach", "Geography", "Years Exp"]
    st.dataframe(display_df, width='stretch', hide_index=True)

    st.markdown('<h2 class="section-header"> Detailed Candidate Profiles</h2>', unsafe_allow_html=True)
    auto_expand = bool(search_query and search_query.strip() and len(filtered) <= 5)

    for idx, row in filtered.iterrows():
        candidate_id = row['id']
        current_roles = st.session_state.flagged_candidates.get(candidate_id, [])

        # Show match info if this is from a search
        match_badges = ""
        if search_query and search_query.strip() and 'match_info' in row and row['match_info']:
            match_badges = " • " + " • ".join(row['match_info'][:3])

        with st.expander(f" {row['name']} — {row['current_title']} at {row['current_company']}", expanded=auto_expand):
            # Create a nice header card
            col_left, col_right = st.columns([2, 1])

            with col_left:
                st.markdown(f"### {row['name']}")
                st.markdown(f"**{row['current_title']}** at **{row['current_company']}**")

                # Show what matched for this candidate
                if search_query and search_query.strip() and 'match_info' in row and row['match_info']:
                    st.markdown("** Matched on:**")
                    for match in row['match_info']:
                        st.markdown(f"- {match}")
                    st.markdown("---")

                if row.get('summary_blurb'):
                    st.markdown(f"*{row['summary_blurb']}*")

            with col_right:
                st.metric("Experience", f"{row['years_experience'] or 0} years")

            st.markdown("---")

            # Key details with badges
            st.markdown("####  Key Details")
            details_col1, details_col2, details_col3 = st.columns(3)

            with details_col1:
                st.markdown(f"** Geography**")
                st.markdown(f'<span class="badge badge-primary">{row["primary_geography"] or "N/A"}</span>', unsafe_allow_html=True)

            with details_col2:
                st.markdown(f"** Sector**")
                st.markdown(f'<span class="badge badge-success">{row["primary_sector"] or "N/A"}</span>', unsafe_allow_html=True)

            with details_col3:
                st.markdown(f"** Investment Approach**")
                st.markdown(f'<span class="badge badge-info">{row["investment_approach"] or "N/A"}</span>', unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Detailed Experience Information
            st.markdown("#### Detailed Experience")
            conn_exp = sqlite3.connect(DB_PATH)
            experiences_df = pd.read_sql_query(
                "SELECT * FROM experiences WHERE candidate_id = ? ORDER BY start_date DESC",
                conn_exp,
                params=(row['id'],)
            )
            conn_exp.close()

            if not experiences_df.empty:
                for _, exp in experiences_df.iterrows():
                    with st.expander(f"**{exp['company']}** - {exp['title']}"):
                        exp_col1, exp_col2 = st.columns(2)

                        with exp_col1:
                            if exp.get('sectors'):
                                try:
                                    sectors = json.loads(exp['sectors']) if isinstance(exp['sectors'], str) else exp['sectors']
                                    if sectors:
                                        st.markdown("**Sectors Covered:**")
                                        sectors_html = ''.join([f'<span class="badge badge-success">{s}</span>' for s in sectors])
                                        st.markdown(sectors_html, unsafe_allow_html=True)
                                except:
                                    pass

                            if exp.get('client_type'):
                                st.markdown(f"**Client Type:** {exp['client_type']}")

                            if exp.get('num_companies_covered'):
                                st.markdown(f"**Companies Covered:** {exp['num_companies_covered']}")

                            if exp.get('coverage_value'):
                                st.markdown(f"**Coverage/AUM:** {exp['coverage_value']}")

                        with exp_col2:
                            if exp.get('regions_covered'):
                                try:
                                    regions = json.loads(exp['regions_covered']) if isinstance(exp['regions_covered'], str) else exp['regions_covered']
                                    if regions:
                                        st.markdown("**Regions Covered:**")
                                        regions_html = ''.join([f'<span class="badge badge-info">{r}</span>' for r in regions])
                                        st.markdown(regions_html, unsafe_allow_html=True)
                                except:
                                    pass

                            if exp.get('sharpe_ratio'):
                                st.markdown(f"**Sharpe Ratio:** {exp['sharpe_ratio']}")

                            if exp.get('alpha'):
                                st.markdown(f"**Alpha:** {exp['alpha']}")

                        if exp.get('valuation_methods_used'):
                            try:
                                methods = json.loads(exp['valuation_methods_used']) if isinstance(exp['valuation_methods_used'], str) else exp['valuation_methods_used']
                                if methods:
                                    st.markdown("**Valuation Methods:**")
                                    methods_html = ''.join([f'<span class="badge badge-primary">{m}</span>' for m in methods])
                                    st.markdown(methods_html, unsafe_allow_html=True)
                            except:
                                pass

                        if exp.get('quant_tools_used'):
                            try:
                                tools = json.loads(exp['quant_tools_used']) if isinstance(exp['quant_tools_used'], str) else exp['quant_tools_used']
                                if tools:
                                    st.markdown("**Quant Tools:**")
                                    tools_html = ''.join([f'<span class="badge badge-info">{t}</span>' for t in tools])
                                    st.markdown(tools_html, unsafe_allow_html=True)
                            except:
                                pass

            # Skills section with badges
            if row.get("all_skills"):
                st.markdown("####  Skills")
                skills_list = str(row['all_skills']).split(',')
                skills_html = ''.join([f'<span class="badge badge-primary">{skill.strip()}</span>' for skill in skills_list[:15]])
                st.markdown(skills_html, unsafe_allow_html=True)
                if len(skills_list) > 15:
                    st.caption(f"+ {len(skills_list) - 15} more skills")

            # Experience section
            if row.get("all_companies"):
                st.markdown("####  Past Companies")
                companies_list = str(row['all_companies']).split(',')
                companies_html = ''.join([f'<span class="badge badge-info">{company.strip()}</span>' for company in companies_list[:10]])
                st.markdown(companies_html, unsafe_allow_html=True)
                if len(companies_list) > 10:
                    st.caption(f"+ {len(companies_list) - 10} more companies")

            # Education section
            if row.get("all_schools"):
                st.markdown("####  Education")
                schools_list = str(row['all_schools']).split(',')
                schools_html = ''.join([f'<span class="badge badge-success">{school.strip()}</span>' for school in schools_list])
                st.markdown(schools_html, unsafe_allow_html=True)

            # Certifications section
            if row.get("certifications"):
                try:
                    certifications = json.loads(row['certifications']) if isinstance(row['certifications'], str) else row['certifications']
                    if certifications:
                        st.markdown("####  Certifications")
                        certs_html = ''.join([f'<span class="badge badge-info">{cert}</span>' for cert in certifications])
                        st.markdown(certs_html, unsafe_allow_html=True)
                except:
                    pass

            st.write("---")

            if row.get("resume_path") and pathlib.Path(row["resume_path"]).exists():
                resume_path = pathlib.Path(row["resume_path"])
                ext = resume_path.suffix.lower()

                with st.expander(" View Resume"):
                    # --- 1 If it's already a PDF ---
                    if ext == ".pdf":
                        with open(resume_path, "rb") as f:
                            base64_pdf = base64.b64encode(f.read()).decode("utf-8")
                        pdf_display = f"""
                            <iframe src="data:application/pdf;base64,{base64_pdf}"
                                    width="100%" height="800px" type="application/pdf">
                            </iframe>
                        """
                        st.markdown(pdf_display, unsafe_allow_html=True)

                    # --- 2 If it's a DOCX, display as formatted text ---
                    elif ext in [".docx", ".doc"]:
                        try:
                            doc = Document(resume_path)

                            # Extract text with basic formatting
                            resume_html = '<div style="background-color: white; padding: 20px; border-radius: 5px; color: black;">'

                            for paragraph in doc.paragraphs:
                                if paragraph.text.strip():
                                    # Detect headings by font size or bold
                                    is_bold = any(run.bold for run in paragraph.runs if run.text.strip())
                                    font_size = max((run.font.size.pt if run.font.size else 11) for run in paragraph.runs if run.text.strip()) if paragraph.runs else 11

                                    if font_size > 13 or (is_bold and len(paragraph.text) < 100):
                                        resume_html += f'<h3 style="color: #1f77b4; margin-top: 15px;">{paragraph.text}</h3>'
                                    else:
                                        resume_html += f'<p style="margin: 5px 0;">{paragraph.text}</p>'

                            # Add tables if any
                            for table in doc.tables:
                                resume_html += '<table style="width: 100%; border-collapse: collapse; margin: 10px 0;">'
                                for row_table in table.rows:
                                    resume_html += '<tr>'
                                    for cell in row_table.cells:
                                        resume_html += f'<td style="border: 1px solid #ddd; padding: 8px;">{cell.text}</td>'
                                    resume_html += '</tr>'
                                resume_html += '</table>'

                            resume_html += '</div>'

                            st.markdown(resume_html, unsafe_allow_html=True)

                            # Also provide download button
                            with open(resume_path, "rb") as f:
                                st.download_button(
                                    label=" Download DOCX",
                                    data=f.read(),
                                    file_name=resume_path.name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                        except Exception as e:
                            st.error(f" Could not display DOCX: {e}")
                            # Fallback: just provide download button
                            with open(resume_path, "rb") as f:
                                st.download_button(
                                    label=" Download Resume",
                                    data=f.read(),
                                    file_name=resume_path.name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )

                    else:
                        st.info(f"Unsupported file type: {ext}")
            else:
                st.info("No resume file available for this candidate.")

            # Flag for roles section - at the bottom after resume
            st.markdown("---")
            st.markdown("####  Flag for Applicable Roles")

            selected_roles = st.multiselect(
                "Select roles this candidate is suitable for",
                options=AVAILABLE_ROLES,
                default=current_roles,
                key=f"roles_{candidate_id}",
                help="Flag this candidate for one or more job roles"
            )

            # Auto-save when selection changes
            if selected_roles != current_roles:
                if selected_roles:
                    st.session_state.flagged_candidates[candidate_id] = selected_roles
                else:
                    # Remove if no roles selected
                    if candidate_id in st.session_state.flagged_candidates:
                        del st.session_state.flagged_candidates[candidate_id]

    # Footer
    st.markdown("<br><br>", unsafe_allow_html=True)

